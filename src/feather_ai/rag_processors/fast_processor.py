"""
This defines an intelligent and smart document processor that takes in
documents of any type and returns chunks (image or text) with bounding box coordinates
for embedding into a vector database
"""
from __future__ import annotations

import asyncio
import os
import re
from concurrent.futures import ProcessPoolExecutor
from typing import List, Dict, Any, Optional, Tuple, Union
import pymupdf
from docx import Document as DocxDocument
import io

from ._doc_convert import _to_pdf
from .. import Document
from google.cloud import vision

# --- Configuration ---
CHUNK_TARGET_SIZE = 1024
MERGE_TOLERANCE = 40

# Quality thresholds
MIN_CHUNK_CHARS = 150
MIN_CHUNK_WORDS = 20
MIN_ALPHA_RATIO = 0.5
MIN_UNIQUE_WORDS_RATIO = 0.3
MIN_AVG_WORD_LENGTH = 2.5
MAX_REPEAT_RATIO = 0.3

# OCR settings
OCR_TRIGGER_MIN_CHARS = 100
OCR_ZOOM_MATRIX = 2.0

# Image filtering
MIN_IMG_WIDTH = 150
MIN_IMG_HEIGHT = 150
MIN_IMG_BYTES = 2048

# Image merging: raw embedded image objects within this many PDF points of each
# other are treated as fragments of the same figure and rendered as one crop.
# The vertical tolerance is larger than the horizontal one: a figure's caption
# strip or axis row typically sits 13-16pt below its plot area (and must merge),
# while side-by-side but unrelated panels/columns sit ~11pt apart horizontally
# (and must stay separate).
IMAGE_MERGE_GAP_X = 12.0
IMAGE_MERGE_GAP_Y = 18.0
IMAGE_RENDER_ZOOM = 2.0
IMAGE_JPEG_QUALITY = 90


# --- Validation Functions ---

def _is_meaningful_chunk(text: str) -> bool:
    """Validates whether a text chunk is meaningful for RAG systems."""
    text = text.strip()

    if len(text) < MIN_CHUNK_CHARS:
        return False

    words = text.split()
    word_count = len(words)

    if word_count < MIN_CHUNK_WORDS:
        return False

    avg_word_len = sum(len(w) for w in words) / word_count if word_count else 0
    if avg_word_len < MIN_AVG_WORD_LENGTH:
        return False

    alpha_chars = sum(1 for c in text if c.isalpha())
    if alpha_chars / len(text) < MIN_ALPHA_RATIO:
        return False

    unique_words = set(w.lower() for w in words if len(w) > 2)
    if len(unique_words) / word_count < MIN_UNIQUE_WORDS_RATIO:
        return False

    repeat_matches = re.findall(r'(.)\1{3,}', text)
    repeat_chars = sum(len(m) + 3 for m in repeat_matches)
    if repeat_chars / len(text) > MAX_REPEAT_RATIO:
        return False

    content_chars = len(text.replace(" ", "").replace("\n", "").replace("\t", ""))
    if content_chars / len(text) < 0.5:
        return False

    return True


def _clean_text(text: str) -> str:
    """Cleans text for embedding quality."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[.]{3,}', '...', text)
    text = re.sub(r'[-]{3,}', ' ', text)
    text = re.sub(r'[_]{3,}', ' ', text)
    text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
    return text.strip()


# --- Extraction Functions ---

def _extract_native_text(
        raw_blocks: List[Tuple],
        page_num: int
) -> List[Dict[str, Any]]:
    """Extracts and chunks text from native PDF text blocks."""
    chunks = []
    current_text = []
    current_boxes = []
    current_bbox = list(raw_blocks[0][:4]) if raw_blocks else [0, 0, 0, 0]

    for block in raw_blocks:
        if block[6] != 0:  # Skip non-text blocks
            continue

        text = block[4].strip()
        if not text:
            continue

        b_x0, b_y0, b_x1, b_y1 = block[:4]
        vertical_gap = b_y0 - current_bbox[3]

        is_new_section = vertical_gap > MERGE_TOLERANCE and current_text
        current_len = sum(len(t) for t in current_text)
        is_chunk_full = current_len > CHUNK_TARGET_SIZE

        if is_new_section or is_chunk_full:
            chunk = _finalize_text_chunk(current_text, current_boxes)
            if chunk:
                chunks.append(chunk)

            current_text = [text]
            current_boxes = [{"page": page_num, "bbox": [b_x0, b_y0, b_x1, b_y1]}]
            current_bbox = [b_x0, b_y0, b_x1, b_y1]
        else:
            current_text.append(text)
            current_boxes.append({"page": page_num, "bbox": [b_x0, b_y0, b_x1, b_y1]})
            current_bbox[0] = min(current_bbox[0], b_x0)
            current_bbox[1] = min(current_bbox[1], b_y0)
            current_bbox[2] = max(current_bbox[2], b_x1)
            current_bbox[3] = max(current_bbox[3], b_y1)

    # Finalize remaining text
    if current_text:
        chunk = _finalize_text_chunk(current_text, current_boxes)
        if chunk:
            chunks.append(chunk)

    return chunks


def _finalize_text_chunk(
        text_parts: List[str],
        boxes: List[Dict[str, Any]]
) -> Optional[Dict[str, Any]]:
    """Validates and finalizes a text chunk."""
    full_text = _clean_text(" ".join(text_parts))

    if not _is_meaningful_chunk(full_text):
        return None

    return {
        "type": "text",
        "content": full_text,
        "boxes": boxes
    }


def _extract_via_ocr(
        page: pymupdf.Page,
        page_num: int,
        vision_client: Any
) -> List[Dict[str, Any]]:
    """Extracts text via Google Cloud Vision OCR."""
    mat = pymupdf.Matrix(OCR_ZOOM_MATRIX, OCR_ZOOM_MATRIX)
    pix = page.get_pixmap(matrix=mat)
    img_bytes = pix.tobytes("png")

    scale_x = pix.width / page.rect.width
    scale_y = pix.height / page.rect.height

    try:
        image = vision.Image(content=img_bytes)
        response = vision_client.document_text_detection(image=image)

        if response.error.message:
            return []

        chunks = []
        for ocr_page in response.full_text_annotation.pages:
            for block in ocr_page.blocks:
                block_text = ""
                for paragraph in block.paragraphs:
                    for word in paragraph.words:
                        word_text = "".join(s.text for s in word.symbols)
                        block_text += word_text + " "

                block_text = _clean_text(block_text)
                if not _is_meaningful_chunk(block_text):
                    continue

                vertices = block.bounding_box.vertices
                xs = [v.x for v in vertices]
                ys = [v.y for v in vertices]

                if not xs or not ys:
                    continue

                bbox = [
                    min(xs) / scale_x,
                    min(ys) / scale_y,
                    max(xs) / scale_x,
                    max(ys) / scale_y
                ]

                chunks.append({
                    "type": "text",
                    "content": block_text,
                    "boxes": [{"page": page_num, "bbox": bbox}]
                })

        return chunks

    except Exception:
        return []


def _bboxes_overlap_or_close(
        a: List[float], b: List[float], gap_x: float, gap_y: float
) -> bool:
    """True if two axis-aligned boxes overlap or are within the gap tolerances."""
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    return not (
        ax1 + gap_x < bx0 or bx1 + gap_x < ax0 or
        ay1 + gap_y < by0 or by1 + gap_y < ay0
    )


def _cluster_bboxes(
        bboxes: List[List[float]], gap_x: float, gap_y: float
) -> List[List[int]]:
    """
    Groups bbox indices into clusters of overlapping/adjacent boxes (union-find).

    PDF producers frequently split a single visual figure into several raster
    XObjects (soft masks, per-channel layers, tiled fragments) stacked on top of
    or right next to each other. Treating each as its own image chunk floods
    downstream consumers with dozens of near-duplicate entries for what is
    visually one diagram, so these are merged before becoming chunks.
    """
    n = len(bboxes)
    parent = list(range(n))

    def find(x: int) -> int:
        while parent[x] != x:
            parent[x] = parent[parent[x]]
            x = parent[x]
        return x

    def union(x: int, y: int) -> None:
        rx, ry = find(x), find(y)
        if rx != ry:
            parent[rx] = ry

    for i in range(n):
        for j in range(i + 1, n):
            if _bboxes_overlap_or_close(bboxes[i], bboxes[j], gap_x, gap_y):
                union(i, j)

    clusters: Dict[int, List[int]] = {}
    for i in range(n):
        clusters.setdefault(find(i), []).append(i)

    return list(clusters.values())


def _union_bbox(bboxes: List[List[float]]) -> List[float]:
    return [
        min(b[0] for b in bboxes),
        min(b[1] for b in bboxes),
        max(b[2] for b in bboxes),
        max(b[3] for b in bboxes),
    ]


def _extract_images(
        doc: pymupdf.Document,
        page: pymupdf.Page,
        page_num: int
) -> List[Dict[str, Any]]:
    """
    Extracts meaningful images from a page.

    Raw embedded image objects that overlap or sit right next to each other are
    merged into a single flattened crop rendered straight from the page, so one
    visual figure yields one chunk regardless of how many raster XObjects the
    PDF producer split it into.
    """
    candidate_bboxes: List[List[float]] = []
    page_area = abs(page.rect) or 1.0

    for img_info in page.get_images(full=True):
        try:
            xref = img_info[0]
            width, height = img_info[2], img_info[3]

            if width < MIN_IMG_WIDTH or height < MIN_IMG_HEIGHT:
                continue

            img_dict = doc.extract_image(xref)
            if len(img_dict["image"]) < MIN_IMG_BYTES:
                continue

            # One XObject may be placed several times (e.g. a frame image
            # reused across a filmstrip) — every placement is a fragment.
            for rect in page.get_image_rects(xref):
                bbox = list(rect)
                if bbox[2] <= bbox[0] or bbox[3] <= bbox[1]:
                    continue
                # Near-full-page rects are backgrounds/watermarks; letting one
                # into the clustering would union the whole page into a single
                # "figure".
                if abs(rect) / page_area > 0.8:
                    continue
                candidate_bboxes.append(bbox)

        except Exception:
            continue

    if not candidate_bboxes:
        return []

    images = []
    for cluster in _cluster_bboxes(candidate_bboxes, IMAGE_MERGE_GAP_X, IMAGE_MERGE_GAP_Y):
        merged_bbox = _union_bbox([candidate_bboxes[i] for i in cluster])

        # Skip crops too small to be a meaningful standalone figure.
        if (merged_bbox[2] - merged_bbox[0] < 40.0
                or merged_bbox[3] - merged_bbox[1] < 30.0):
            continue

        try:
            pix = page.get_pixmap(
                matrix=pymupdf.Matrix(IMAGE_RENDER_ZOOM, IMAGE_RENDER_ZOOM),
                clip=pymupdf.Rect(*merged_bbox),
            )
            img_bytes = pix.tobytes("jpg", jpg_quality=IMAGE_JPEG_QUALITY)
        except Exception:
            continue

        if len(img_bytes) < MIN_IMG_BYTES:
            continue

        images.append({
            "type": "image",
            "content": img_bytes,
            "boxes": [{"page": page_num, "bbox": merged_bbox}]
        })

    return images


def _process_docx_document(document: Document) -> List[Dict[str, Any]]:
    """
    Processes a DOCX document by extracting text and chunking it.
    Simple text-only extraction without images.
    """
    try:
        docx_file = DocxDocument(io.BytesIO(document.content))
    except Exception as e:
        print(f"Error opening DOCX: {e}")
        return []

    all_chunks = []
    current_text = []
    current_char_count = 0
    total_paragraphs = 0
    non_empty_paragraphs = 0

    for paragraph in docx_file.paragraphs:
        total_paragraphs += 1
        text = paragraph.text.strip()
        if not text:
            continue

        non_empty_paragraphs += 1
        paragraph_len = len(text)

        # Check if adding this paragraph would exceed chunk size
        if current_char_count + paragraph_len > CHUNK_TARGET_SIZE and current_text:
            # Finalize current chunk
            full_text = _clean_text(" ".join(current_text))
            is_meaningful = _is_meaningful_chunk(full_text)
            if not is_meaningful:
                print(f"Chunk rejected: {len(full_text)} chars, {len(full_text.split())} words")
                print(f"First 100 chars: {full_text[:100]}")
            if is_meaningful:
                all_chunks.append({
                    "type": "text",
                    "content": full_text,
                    "boxes": [{"page": 0, "bbox": [0, 0, 0, 0]}]  # No bbox for DOCX
                })

            current_text = []
            current_char_count = 0

        current_text.append(text)
        current_char_count += paragraph_len

    # Finalize remaining text
    if current_text:
        full_text = _clean_text(" ".join(current_text))
        is_meaningful = _is_meaningful_chunk(full_text)
        if not is_meaningful:
            print(f"Final chunk rejected: {len(full_text)} chars, {len(full_text.split())} words")
            print(f"First 100 chars: {full_text[:100]}")
        if is_meaningful:
            all_chunks.append({
                "type": "text",
                "content": full_text,
                "boxes": [{"page": 0, "bbox": [0, 0, 0, 0]}]
            })

    return all_chunks


# --- Document Processing ---

def _process_single_document(doc_tuple: Tuple[Document, str]) -> List[Dict[str, Any]]:
    """
    Processes a single document synchronously.
    Designed to run in a process pool worker.
    Routes to appropriate processor based on document type.
    """
    document, doc_type = doc_tuple

    if doc_type == 'docx':
        return _process_docx_document(document)

    # PDF processing
    doc = pymupdf.Document(stream=document.content)
    vision_client = None
    all_chunks = []

    for page_num in range(len(doc)):
        page = doc[page_num]

        raw_blocks = page.get_text("blocks", sort=True)
        native_text_len = sum(
            len(b[4].strip()) for b in raw_blocks if b[6] == 0
        )

        if native_text_len < OCR_TRIGGER_MIN_CHARS:
            if vision_client is None:
                vision_client = vision.ImageAnnotatorClient()
            text_chunks = _extract_via_ocr(page, page_num, vision_client)

            all_chunks.extend(text_chunks)
        else:
            text_chunks = _extract_native_text(raw_blocks, page_num)
            all_chunks.extend(text_chunks)

            image_chunks = _extract_images(doc, page, page_num)
            all_chunks.extend(image_chunks)

    doc.close()
    return all_chunks


# --- Public API ---

class FastProcessor:
    """
    Async document processor optimized for RAG pipelines.
    Supports PDF and DOCX files.

    Example:
        processor = FastProcessor()
        results = await processor.process_documents([doc1, doc2])

        # Or synchronously:
        results = processor.process_documents_sync([doc1, doc2])
    """

    def __init__(self, max_workers: Optional[int] = None):
        """
        Initialize the processor.

        Args:
            max_workers: Maximum parallel workers. Defaults to min(4, cpu_count).
        """
        self.max_workers = max_workers or min(4, os.cpu_count() or 1)
        pymupdf.TOOLS.store_shrink(75)

    async def process_documents(
            self,
            documents: List[Union[str, Document]]
    ) -> List[List[Dict[str, Any]]]:
        """
        Process multiple documents concurrently.

        Args:
            documents: List of file paths or Document objects.

        Returns:
            List of chunk lists (one per document).
            Each chunk dict contains:
                - type: "text" | "image"
                - content: str (for text) or bytes (for image)
                - boxes: List of {"page": int, "bbox": [x0, y0, x1, y1]}
        """
        normalized = self._normalize_inputs(documents)

        loop = asyncio.get_running_loop()

        with ProcessPoolExecutor(max_workers=self.max_workers) as executor:
            tasks = [
                loop.run_in_executor(executor, _process_single_document, doc)
                for doc in normalized
            ]
            results = await asyncio.gather(*tasks)

        return list(results)

    def process_documents_sync(
            self,
            documents: List[Union[str, Document]]
    ) -> List[List[Dict[str, Any]]]:
        """
        Synchronous wrapper for process_documents.

        Useful when not in an async context.
        """
        return asyncio.run(self.process_documents(documents))

    def _normalize_inputs(
            self,
            documents: List[Union[str, Document]]
    ) -> List[Tuple[Document, str]]:
        """Converts file paths to Document objects, returns (Document, type) tuples"""
        normalized = []
        for doc in documents:
            if isinstance(doc, str):
                doc_object = Document.from_path(doc)
            else:
                doc_object = doc

            # Check if this is a DOCX file by mime type or extension
            is_docx = (
                doc_object.mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                or doc_object.filename.lower().endswith('.docx')
            )

            if is_docx:
                # Keep DOCX as-is, don't convert to PDF
                normalized.append((doc_object, 'docx'))
            else:
                # Convert other document types to PDF
                pdf_doc = _to_pdf(doc_object)
                normalized.append((pdf_doc, 'pdf'))

        return normalized


async def chunk_documents(
        documents: List[Union[str, Document]],
        max_workers: Optional[int] = None
) -> List[List[Dict[str, Any]]]:
    """
    Convenience function to process documents without instantiating FastProcessor.

    Args:
        documents: List of file paths or Document objects.
        max_workers: Maximum parallel workers.

    Returns:
        List of chunk lists. See FastProcessor.process_documents for format.

    Example:
        chunks = await chunk_documents(["doc1.pdf", "doc2.pdf"])
        for doc_chunks in chunks:
            for chunk in doc_chunks:
                if chunk["type"] == "text":
                    embed(chunk["content"])
    """
    processor = FastProcessor(max_workers=max_workers)
    return await processor.process_documents(documents)